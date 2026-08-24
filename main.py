import os
import io
import json
import base64
import time
import urllib.request
import asyncio
import threading
from http.server import HTTPServer, BaseHTTPRequestHandler

from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------
# 1. خادم وهمي لإبقاء الخدمة نشطة على Render (Health Check)
# ---------------------------------------------------------
class HealthCheckHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.end_headers()
        self.wfile.write(b"OK")

def run_web_port():
    port = int(os.environ.get("PORT", 10000))
    server = HTTPServer(('0.0.0.0', port), HealthCheckHandler)
    server.serve_forever()

threading.Thread(target=run_web_port, daemon=True).start()

# ---------------------------------------------------------
# 2. إعداد المفاتيح والدوال
# ---------------------------------------------------------
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

def call_gemini_direct(prompt: str, image_bytes: bytes = None) -> dict:
    if not GEMINI_API_KEY:
        raise Exception("مفتاح GEMINI_API_KEY غير موجود في متغيرات البيئة!")

    # جلب قائمة النماذج المتاحة للخدمة
    list_url = f"https://generativelanguage.googleapis.com/v1beta/models?key={GEMINI_API_KEY}"
    req_list = urllib.request.Request(list_url)
    
    try:
        with urllib.request.urlopen(req_list, timeout=15) as res:
            models_data = json.loads(res.read().decode('utf-8'))
            
        valid_models = []
        for m in models_data.get("models", []):
            methods = m.get("supportedGenerationMethods", [])
            if "generateContent" in methods:
                valid_models.append(m['name'])
    except Exception as e:
        raise Exception(f"خطأ في الاتصال بخدمة جوجل: {e}")

    if not valid_models:
        raise Exception("لم يتم العثور على أي نموذج مفعل متاح لك.")

    # توجيهات الذكاء الاصطناعي لاستخراج بيانات المستند
    system_instruction = (
        "أنت خبير في تحليل الوثائق والصور وتصميم مستندات Word. قم بتحليل الطلب والصورة المرفقة (إن وجدت) "
        "وأرجع الناتج بصيغة JSON فقط بدون أسلوب markdown:\n"
        "{\n"
        '  "header_top": "الجمهورية الجزائرية الديمقراطية الشعبية",\n'
        '  "title": "مذكرة إدماج جزئي",\n'
        '  "metadata": ["الميدان: ...", "المقطع التعلمي: ...", "المستوى: ..."],\n'
        '  "has_table": true,\n'
        '  "table_columns": ["نص الوضعية", "الحل", "المورد المستهدف"],\n'
        '  "rows_count": 5,\n'
        '  "footer_notes": ""\n'
        "}"
    )

    parts = []
    # تحويل الصورة إلى base64 إذا كانت متوفرة
    if image_bytes:
        b64_image = base64.b64encode(image_bytes).decode('utf-8')
        parts.append({
            "inlineData": {
                "mimeType": "image/jpeg",
                "data": b64_image
            }
        })

    full_text = f"{system_instruction}\n\nطلب المستخدم:\n{prompt if prompt else 'قم بإنشاء وثيقة مطابقة لهذه الصورة'}"
    parts.append({"text": full_text})

    payload = {"contents": [{"parts": parts}]}
    data = json.dumps(payload).encode('utf-8')

    last_err = None
    for model_full_name in valid_models:
        url = f"https://generativelanguage.googleapis.com/v1beta/{model_full_name}:generateContent?key={GEMINI_API_KEY}"
        try:
            req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json'})
            with urllib.request.urlopen(req, timeout=30) as response:
                res_body = response.read().decode('utf-8')
                res_json = json.loads(res_body)
                text = res_json['candidates'][0]['content']['parts'][0]['text'].strip()
                
                # تنظيف النص من أوسمة البرمجة
                if text.startswith("```json"): text = text[7:]
                if text.startswith("```"): text = text[3:]
                if text.endswith("```"): text = text[:-3]
                return json.loads(text.strip())
        except Exception as e:
            last_err = e
            continue

    raise Exception(f"تعذر معالجة البيانات: {last_err}")

# ---------------------------------------------------------
# 3. دالة بناء ملف Word (python-docx)
# ---------------------------------------------------------
def build_docx(doc_data: dict) -> io.BytesIO:
    doc = Document()
    
    # ضبط الهوامش
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    # الترويسة العلوية
    if doc_data.get("header_top"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(doc_data["header_top"])
        r.font.bold = True
        r.font.size = Pt(11)

    # العنوان الرئيسي
    if doc_data.get("title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(doc_data["title"])
        r.font.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0, 51, 102)

    # المعلومات الفرعية (البيانات)
    if doc_data.get("metadata"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("  |  ".join(doc_data["metadata"]))
        r.font.bold = True
        r.font.size = Pt(10)

    # إضافة الجدول
    if doc_data.get("has_table", False):
        cols = doc_data.get("table_columns", ["العمود 1", "العمود 2"])
        rows_cnt = doc_data.get("rows_count", 5)
        
        table = doc.add_table(rows=rows_cnt + 1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for idx, col_name in enumerate(cols):
            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = hdr_cells[idx].paragraphs[0].add_run(col_name)
            r.font.bold = True
            r.font.size = Pt(10)

    # الهامش السفلي
    if doc_data.get("footer_notes"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(doc_data["footer_notes"])
        r.font.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# ---------------------------------------------------------
# 4. معالجات أوامر ورسائل تليغرام
# ---------------------------------------------------------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("مرحباً بك! يمكنك إرسال نصوص أو صور لمستندات وسأقوم بتحليلها وإنشاء ملف Word مطابِق لك.")

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    status = await update.message.reply_text("⏳ جاري تحليل الطلب والصورة وبناء المستند...")
    try:
        prompt_text = update.message.text or update.message.caption or ""
        image_bytes = None

        # استخراج الصورة في حال إرسال صورة
        if update.message.photo:
            photo_file = await update.message.photo[-1].get_file()
            image_bytes = await photo_file.download_as_bytearray()

        # معالجة الطلب في الخلفية
        data = await asyncio.to_thread(call_gemini_direct, prompt_text, image_bytes)
        doc_bytes = await asyncio.to_thread(build_docx, data)
        
        await update.message.reply_document(
            document=doc_bytes, 
            filename="Document.docx", 
            caption="✨ تم إنشاء المستند بنجاح!"
        )
        await status.delete()
    except Exception as e:
        await status.edit_text(f"❌ حدث خطأ: {str(e)[:250]}")

# ---------------------------------------------------------
# 5. التشغيل الرئيسي مع تفادي تعارض Polling
# ---------------------------------------------------------
if __name__ == '__main__':
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()
    
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler((filters.TEXT | filters.PHOTO) & ~filters.COMMAND, handle_message))

    print("🤖 البوت يعمل الآن...")

    # حلقة تكرار لإعادة الاتصال التلقائي في حال وجود تعارض مؤقت أثناء Deployment على Render
    while True:
        try:
            app.run_polling(drop_pending_updates=True, stop_signals=None)
            break
        except Exception as e:
            print(f"⚠️ تعارض أو خطأ شبكة مؤقت: {e}. إعادة المحاولة بعد 10 ثوانٍ...")
            time.sleep(10)
